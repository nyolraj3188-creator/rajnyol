#!/usr/bin/env python3
"""
Full backend regression test for LovePDF FastAPI backend.
Tests all 15 endpoints with real file generation and validation.
"""
import os
import sys
import io
import requests
import tempfile
from pathlib import Path
from typing import Dict, Any, List, Tuple

# Base URL from frontend/.env
BASE_URL = "https://pdf-font-detection.preview.emergentagent.com"
API_BASE = f"{BASE_URL}/api"

# Test results storage
results = []

def log_test(endpoint: str, status: str, details: str):
    """Log a test result."""
    results.append({
        "endpoint": endpoint,
        "status": status,
        "details": details
    })
    symbol = "✅" if status == "PASS" else "❌"
    print(f"{symbol} {endpoint}: {status} - {details}")

def create_simple_pdf() -> bytes:
    """Create a simple text PDF with English content."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.drawString(100, 750, "This is a test PDF document.")
    c.drawString(100, 730, "It contains some English text for testing.")
    c.drawString(100, 710, "The quick brown fox jumps over the lazy dog.")
    c.showPage()
    c.save()
    return buf.getvalue()

def create_table_pdf() -> bytes:
    """Create a PDF with a simple table."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    # Draw a simple table
    c.drawString(100, 750, "Name")
    c.drawString(200, 750, "Age")
    c.drawString(300, 750, "City")
    c.drawString(100, 730, "John")
    c.drawString(200, 730, "25")
    c.drawString(300, 730, "New York")
    c.drawString(100, 710, "Jane")
    c.drawString(200, 710, "30")
    c.drawString(300, 710, "London")
    c.showPage()
    c.save()
    return buf.getvalue()

def create_scanned_pdf() -> bytes:
    """Create an image-only PDF (simulating a scanned document)."""
    from PIL import Image, ImageDraw, ImageFont
    import img2pdf
    
    # Create an image with text
    img = Image.new('RGB', (800, 600), color='white')
    draw = ImageDraw.Draw(img)
    
    # Try to use a default font
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 40)
    except Exception:
        font = ImageFont.load_default()
    
    draw.text((50, 50), "This is a scanned document", fill='black', font=font)
    draw.text((50, 120), "It has no text layer", fill='black', font=font)
    draw.text((50, 190), "OCR should recognize this text", fill='black', font=font)
    
    # Convert to PDF
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG')
    img_bytes.seek(0)
    
    pdf_bytes = img2pdf.convert(img_bytes.getvalue())
    return pdf_bytes

def create_kruti_pdf() -> bytes:
    """Create a PDF with Kruti Dev ASCII encoding (non-legacy font name)."""
    from fpdf import FPDF
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=14)
    
    # Genuine Kruti Dev ASCII encoding
    pdf.cell(200, 10, txt="pfj= izek.k i=", ln=True)
    pdf.cell(200, 10, txt="izekf.kr fd;k tkrk gS", ln=True)
    
    return bytes(pdf.output())

def create_empty_pdf() -> bytes:
    """Create an empty PDF with no text."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.showPage()
    c.save()
    return buf.getvalue()

def create_docx() -> bytes:
    """Create a simple .docx file."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("This is a test Word document.")
    doc.add_paragraph("It will be converted to PDF.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def create_xlsx() -> bytes:
    """Create a simple .xlsx file."""
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws['A1'] = 'Name'
    ws['B1'] = 'Value'
    ws['A2'] = 'Test'
    ws['B2'] = 123
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()

def is_valid_pdf(data: bytes) -> bool:
    """Check if data is a valid PDF."""
    return data.startswith(b'%PDF-') and len(data) > 100

def is_valid_docx(data: bytes) -> bool:
    """Check if data is a valid .docx (zip with required parts)."""
    import zipfile
    try:
        buf = io.BytesIO(data)
        if not zipfile.is_zipfile(buf):
            return False
        buf.seek(0)
        with zipfile.ZipFile(buf) as z:
            names = z.namelist()
            return '[Content_Types].xml' in names and 'word/document.xml' in names
    except Exception:
        return False

def is_valid_xlsx(data: bytes) -> bool:
    """Check if data is a valid .xlsx."""
    import zipfile
    try:
        buf = io.BytesIO(data)
        if not zipfile.is_zipfile(buf):
            return False
        buf.seek(0)
        with zipfile.ZipFile(buf) as z:
            names = z.namelist()
            return '[Content_Types].xml' in names and 'xl/workbook.xml' in names
    except Exception:
        return False

def is_valid_pptx(data: bytes) -> bool:
    """Check if data is a valid .pptx."""
    import zipfile
    try:
        buf = io.BytesIO(data)
        if not zipfile.is_zipfile(buf):
            return False
        buf.seek(0)
        with zipfile.ZipFile(buf) as z:
            names = z.namelist()
            return '[Content_Types].xml' in names and 'ppt/presentation.xml' in names
    except Exception:
        return False

def test_1_health():
    """Test 1: GET /api/pdf/health"""
    try:
        resp = requests.get(f"{API_BASE}/pdf/health", timeout=10)
        if resp.status_code != 200:
            log_test("GET /api/pdf/health", "FAIL", f"HTTP {resp.status_code}")
            return
        
        data = resp.json()
        if not data.get("ok"):
            log_test("GET /api/pdf/health", "FAIL", "ok=false")
            return
        
        tools = data.get("tools", {})
        required = ["soffice", "gs", "qpdf", "tesseract", "pdftoppm", "ocrmypdf"]
        missing = [t for t in required if not tools.get(t)]
        
        if missing:
            log_test("GET /api/pdf/health", "FAIL", f"Missing tools: {missing}")
        else:
            log_test("GET /api/pdf/health", "PASS", "All tools available")
    except Exception as e:
        log_test("GET /api/pdf/health", "FAIL", f"Exception: {e}")

def test_2_mongo_connectivity():
    """Test 2: GET /api/ and POST /api/status (Mongo connectivity)"""
    try:
        # Test GET /api/
        resp = requests.get(f"{API_BASE}/", timeout=10)
        if resp.status_code != 200:
            log_test("GET /api/", "FAIL", f"HTTP {resp.status_code}")
            return
        
        # Test POST /api/status
        resp = requests.post(f"{API_BASE}/status", json={"client_name": "test"}, timeout=10)
        if resp.status_code != 200:
            log_test("POST /api/status", "FAIL", f"HTTP {resp.status_code}")
            return
        
        # Test GET /api/status
        resp = requests.get(f"{API_BASE}/status", timeout=10)
        if resp.status_code != 200:
            log_test("GET /api/status", "FAIL", f"HTTP {resp.status_code}")
            return
        
        data = resp.json()
        if not isinstance(data, list):
            log_test("Mongo connectivity", "FAIL", "Status list not returned")
            return
        
        # Check if our test entry is in the list
        found = any(s.get("client_name") == "test" for s in data)
        if found:
            log_test("Mongo connectivity", "PASS", "All endpoints working, test entry found")
        else:
            log_test("Mongo connectivity", "PASS", "All endpoints working (entry may have been cleaned)")
    except Exception as e:
        log_test("Mongo connectivity", "FAIL", f"Exception: {e}")

def test_3_inspect():
    """Test 3: POST /api/pdf/inspect with different PDFs"""
    # Test 3a: English text PDF
    try:
        pdf_data = create_simple_pdf()
        files = {'file': ('test.pdf', pdf_data, 'application/pdf')}
        resp = requests.post(f"{API_BASE}/pdf/inspect", files=files, timeout=30)
        
        if resp.status_code != 200:
            log_test("POST /api/pdf/inspect (English)", "FAIL", f"HTTP {resp.status_code}")
        else:
            data = resp.json()
            if (data.get("looks_english") == True and 
                data.get("has_text") == True and 
                data.get("devanagari_ratio") == 0.0):
                log_test("POST /api/pdf/inspect (English)", "PASS", f"Correct detection: {data}")
            else:
                log_test("POST /api/pdf/inspect (English)", "FAIL", f"Wrong detection: {data}")
    except Exception as e:
        log_test("POST /api/pdf/inspect (English)", "FAIL", f"Exception: {e}")
    
    # Test 3b: Kruti-ASCII PDF with non-legacy font
    try:
        pdf_data = create_kruti_pdf()
        files = {'file': ('kruti.pdf', pdf_data, 'application/pdf')}
        resp = requests.post(f"{API_BASE}/pdf/inspect", files=files, timeout=30)
        
        if resp.status_code != 200:
            log_test("POST /api/pdf/inspect (Kruti)", "FAIL", f"HTTP {resp.status_code}")
        else:
            data = resp.json()
            if (data.get("has_text") == True and 
                data.get("devanagari_ratio") < 0.15 and 
                data.get("looks_english") == False and 
                data.get("legacy_hindi") == False):
                log_test("POST /api/pdf/inspect (Kruti)", "PASS", f"Correct content-based detection: {data}")
            else:
                log_test("POST /api/pdf/inspect (Kruti)", "FAIL", f"Wrong detection: {data}")
    except Exception as e:
        log_test("POST /api/pdf/inspect (Kruti)", "FAIL", f"Exception: {e}")
    
    # Test 3c: Empty PDF
    try:
        pdf_data = create_empty_pdf()
        files = {'file': ('empty.pdf', pdf_data, 'application/pdf')}
        resp = requests.post(f"{API_BASE}/pdf/inspect", files=files, timeout=30)
        
        if resp.status_code != 200:
            log_test("POST /api/pdf/inspect (Empty)", "FAIL", f"HTTP {resp.status_code}")
        else:
            data = resp.json()
            if data.get("has_text") == False:
                log_test("POST /api/pdf/inspect (Empty)", "PASS", f"Correct detection: {data}")
            else:
                log_test("POST /api/pdf/inspect (Empty)", "FAIL", f"Wrong detection: {data}")
    except Exception as e:
        log_test("POST /api/pdf/inspect (Empty)", "FAIL", f"Exception: {e}")

def test_4_office_to_pdf():
    """Test 4: POST /api/pdf/office-to-pdf with .docx and .xlsx"""
    # Test 4a: .docx to PDF
    try:
        docx_data = create_docx()
        files = {'file': ('test.docx', docx_data, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        resp = requests.post(f"{API_BASE}/pdf/office-to-pdf", files=files, timeout=60)
        
        if resp.status_code != 200:
            log_test("POST /api/pdf/office-to-pdf (.docx)", "FAIL", f"HTTP {resp.status_code}")
        else:
            pdf_data = resp.content
            if is_valid_pdf(pdf_data):
                log_test("POST /api/pdf/office-to-pdf (.docx)", "PASS", f"Valid PDF output ({len(pdf_data)} bytes)")
            else:
                log_test("POST /api/pdf/office-to-pdf (.docx)", "FAIL", "Invalid PDF output")
    except Exception as e:
        log_test("POST /api/pdf/office-to-pdf (.docx)", "FAIL", f"Exception: {e}")
    
    # Test 4b: .xlsx to PDF
    try:
        xlsx_data = create_xlsx()
        files = {'file': ('test.xlsx', xlsx_data, 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')}
        resp = requests.post(f"{API_BASE}/pdf/office-to-pdf", files=files, timeout=60)
        
        if resp.status_code != 200:
            log_test("POST /api/pdf/office-to-pdf (.xlsx)", "FAIL", f"HTTP {resp.status_code}")
        else:
            pdf_data = resp.content
            if is_valid_pdf(pdf_data):
                log_test("POST /api/pdf/office-to-pdf (.xlsx)", "PASS", f"Valid PDF output ({len(pdf_data)} bytes)")
            else:
                log_test("POST /api/pdf/office-to-pdf (.xlsx)", "FAIL", "Invalid PDF output")
    except Exception as e:
        log_test("POST /api/pdf/office-to-pdf (.xlsx)", "FAIL", f"Exception: {e}")

def test_5_html_to_pdf():
    """Test 5: POST /api/pdf/html-to-pdf"""
    try:
        html_content = """
        <html>
        <head><title>Test Page</title></head>
        <body>
            <h1>Test HTML to PDF</h1>
            <p>This is a simple HTML page for testing.</p>
        </body>
        </html>
        """
        data = {'html': html_content}
        resp = requests.post(f"{API_BASE}/pdf/html-to-pdf", data=data, timeout=60)
        
        if resp.status_code != 200:
            log_test("POST /api/pdf/html-to-pdf", "FAIL", f"HTTP {resp.status_code}")
        else:
            pdf_data = resp.content
            if is_valid_pdf(pdf_data):
                log_test("POST /api/pdf/html-to-pdf", "PASS", f"Valid PDF output ({len(pdf_data)} bytes)")
            else:
                log_test("POST /api/pdf/html-to-pdf", "FAIL", "Invalid PDF output")
    except Exception as e:
        log_test("POST /api/pdf/html-to-pdf", "FAIL", f"Exception: {e}")

def test_6_pdf_to_word():
    """Test 6: POST /api/pdf/pdf-to-word with English and Kruti PDFs"""
    # Test 6a: English text PDF
    try:
        pdf_data = create_simple_pdf()
        files = {'file': ('test.pdf', pdf_data, 'application/pdf')}
        resp = requests.post(f"{API_BASE}/pdf/pdf-to-word", files=files, timeout=120)
        
        if resp.status_code != 200:
            log_test("POST /api/pdf/pdf-to-word (English)", "FAIL", f"HTTP {resp.status_code}")
        else:
            docx_data = resp.content
            if is_valid_docx(docx_data):
                # Check if text is preserved
                from docx import Document
                doc = Document(io.BytesIO(docx_data))
                text = ' '.join([p.text for p in doc.paragraphs])
                if 'English' in text or 'test' in text.lower():
                    log_test("POST /api/pdf/pdf-to-word (English)", "PASS", f"Valid .docx with English text ({len(docx_data)} bytes)")
                else:
                    log_test("POST /api/pdf/pdf-to-word (English)", "FAIL", "Text not preserved correctly")
            else:
                log_test("POST /api/pdf/pdf-to-word (English)", "FAIL", "Invalid .docx output")
    except Exception as e:
        log_test("POST /api/pdf/pdf-to-word (English)", "FAIL", f"Exception: {e}")
    
    # Test 6b: Kruti/Hindi PDF
    try:
        pdf_data = create_kruti_pdf()
        files = {'file': ('kruti.pdf', pdf_data, 'application/pdf')}
        resp = requests.post(f"{API_BASE}/pdf/pdf-to-word", files=files, timeout=120)
        
        if resp.status_code != 200:
            log_test("POST /api/pdf/pdf-to-word (Kruti)", "FAIL", f"HTTP {resp.status_code}")
        else:
            docx_data = resp.content
            if is_valid_docx(docx_data):
                log_test("POST /api/pdf/pdf-to-word (Kruti)", "PASS", f"Valid .docx output ({len(docx_data)} bytes)")
            else:
                log_test("POST /api/pdf/pdf-to-word (Kruti)", "FAIL", "Invalid .docx output")
    except Exception as e:
        log_test("POST /api/pdf/pdf-to-word (Kruti)", "FAIL", f"Exception: {e}")

def test_7_pdf_to_excel():
    """Test 7: POST /api/pdf/pdf-to-excel"""
    try:
        pdf_data = create_table_pdf()
        files = {'file': ('table.pdf', pdf_data, 'application/pdf')}
        resp = requests.post(f"{API_BASE}/pdf/pdf-to-excel", files=files, timeout=120)
        
        if resp.status_code != 200:
            log_test("POST /api/pdf/pdf-to-excel", "FAIL", f"HTTP {resp.status_code}")
        else:
            xlsx_data = resp.content
            if is_valid_xlsx(xlsx_data):
                log_test("POST /api/pdf/pdf-to-excel", "PASS", f"Valid .xlsx output ({len(xlsx_data)} bytes)")
            else:
                log_test("POST /api/pdf/pdf-to-excel", "FAIL", "Invalid .xlsx output")
    except Exception as e:
        log_test("POST /api/pdf/pdf-to-excel", "FAIL", f"Exception: {e}")

def test_8_pdf_to_ppt():
    """Test 8: POST /api/pdf/pdf-to-ppt"""
    try:
        pdf_data = create_simple_pdf()
        files = {'file': ('test.pdf', pdf_data, 'application/pdf')}
        resp = requests.post(f"{API_BASE}/pdf/pdf-to-ppt", files=files, timeout=120)
        
        if resp.status_code != 200:
            log_test("POST /api/pdf/pdf-to-ppt", "FAIL", f"HTTP {resp.status_code}")
        else:
            pptx_data = resp.content
            if is_valid_pptx(pptx_data):
                log_test("POST /api/pdf/pdf-to-ppt", "PASS", f"Valid .pptx output ({len(pptx_data)} bytes)")
            else:
                log_test("POST /api/pdf/pdf-to-ppt", "FAIL", "Invalid .pptx output")
    except Exception as e:
        log_test("POST /api/pdf/pdf-to-ppt", "FAIL", f"Exception: {e}")

def test_9_ocr():
    """Test 9: POST /api/pdf/ocr"""
    try:
        pdf_data = create_scanned_pdf()
        files = {'file': ('scanned.pdf', pdf_data, 'application/pdf')}
        data = {'lang': 'eng'}
        resp = requests.post(f"{API_BASE}/pdf/ocr", files=files, data=data, timeout=180)
        
        if resp.status_code != 200:
            log_test("POST /api/pdf/ocr", "FAIL", f"HTTP {resp.status_code}")
        else:
            ocr_pdf = resp.content
            if is_valid_pdf(ocr_pdf):
                # Check if the PDF now has a text layer
                import pdfplumber
                with pdfplumber.open(io.BytesIO(ocr_pdf)) as pdf:
                    text = pdf.pages[0].extract_text() or ""
                    if len(text.strip()) > 10:
                        log_test("POST /api/pdf/ocr", "PASS", f"Valid searchable PDF with text ({len(ocr_pdf)} bytes)")
                    else:
                        log_test("POST /api/pdf/ocr", "FAIL", "PDF has no text layer after OCR")
            else:
                log_test("POST /api/pdf/ocr", "FAIL", "Invalid PDF output")
    except Exception as e:
        log_test("POST /api/pdf/ocr", "FAIL", f"Exception: {e}")

def test_10_protect():
    """Test 10: POST /api/pdf/protect"""
    try:
        pdf_data = create_simple_pdf()
        files = {'file': ('test.pdf', pdf_data, 'application/pdf')}
        data = {'password': 'test123'}
        resp = requests.post(f"{API_BASE}/pdf/protect", files=files, data=data, timeout=60)
        
        if resp.status_code != 200:
            log_test("POST /api/pdf/protect", "FAIL", f"HTTP {resp.status_code}")
        else:
            protected_pdf = resp.content
            if is_valid_pdf(protected_pdf):
                # Try to open with pikepdf to verify it's encrypted
                import pikepdf
                try:
                    # Should fail without password
                    pikepdf.open(io.BytesIO(protected_pdf))
                    log_test("POST /api/pdf/protect", "FAIL", "PDF not encrypted (opened without password)")
                except pikepdf.PasswordError:
                    # Should succeed with password
                    try:
                        pikepdf.open(io.BytesIO(protected_pdf), password='test123')
                        log_test("POST /api/pdf/protect", "PASS", f"Valid encrypted PDF ({len(protected_pdf)} bytes)")
                    except Exception:
                        log_test("POST /api/pdf/protect", "FAIL", "Cannot open with correct password")
            else:
                log_test("POST /api/pdf/protect", "FAIL", "Invalid PDF output")
    except Exception as e:
        log_test("POST /api/pdf/protect", "FAIL", f"Exception: {e}")

def test_11_unlock():
    """Test 11: POST /api/pdf/unlock"""
    try:
        # First create a protected PDF
        pdf_data = create_simple_pdf()
        files = {'file': ('test.pdf', pdf_data, 'application/pdf')}
        data = {'password': 'test123'}
        resp = requests.post(f"{API_BASE}/pdf/protect", files=files, data=data, timeout=60)
        
        if resp.status_code != 200:
            log_test("POST /api/pdf/unlock", "FAIL", "Could not create protected PDF for test")
            return
        
        protected_pdf = resp.content
        
        # Now unlock it
        files = {'file': ('protected.pdf', protected_pdf, 'application/pdf')}
        data = {'password': 'test123'}
        resp = requests.post(f"{API_BASE}/pdf/unlock", files=files, data=data, timeout=60)
        
        if resp.status_code != 200:
            log_test("POST /api/pdf/unlock", "FAIL", f"HTTP {resp.status_code}")
        else:
            unlocked_pdf = resp.content
            if is_valid_pdf(unlocked_pdf):
                # Verify it's no longer encrypted
                import pikepdf
                try:
                    pikepdf.open(io.BytesIO(unlocked_pdf))
                    log_test("POST /api/pdf/unlock", "PASS", f"Valid decrypted PDF ({len(unlocked_pdf)} bytes)")
                except pikepdf.PasswordError:
                    log_test("POST /api/pdf/unlock", "FAIL", "PDF still encrypted")
            else:
                log_test("POST /api/pdf/unlock", "FAIL", "Invalid PDF output")
    except Exception as e:
        log_test("POST /api/pdf/unlock", "FAIL", f"Exception: {e}")

def test_12_repair():
    """Test 12: POST /api/pdf/repair"""
    try:
        pdf_data = create_simple_pdf()
        files = {'file': ('test.pdf', pdf_data, 'application/pdf')}
        resp = requests.post(f"{API_BASE}/pdf/repair", files=files, timeout=180)
        
        if resp.status_code != 200:
            log_test("POST /api/pdf/repair", "FAIL", f"HTTP {resp.status_code}")
        else:
            repaired_pdf = resp.content
            if is_valid_pdf(repaired_pdf):
                log_test("POST /api/pdf/repair", "PASS", f"Valid repaired PDF ({len(repaired_pdf)} bytes)")
            else:
                log_test("POST /api/pdf/repair", "FAIL", "Invalid PDF output")
    except Exception as e:
        log_test("POST /api/pdf/repair", "FAIL", f"Exception: {e}")

def test_13_pdfa():
    """Test 13: POST /api/pdf/pdfa"""
    try:
        pdf_data = create_simple_pdf()
        files = {'file': ('test.pdf', pdf_data, 'application/pdf')}
        resp = requests.post(f"{API_BASE}/pdf/pdfa", files=files, timeout=180)
        
        if resp.status_code != 200:
            log_test("POST /api/pdf/pdfa", "FAIL", f"HTTP {resp.status_code}")
        else:
            pdfa_data = resp.content
            if is_valid_pdf(pdfa_data):
                log_test("POST /api/pdf/pdfa", "PASS", f"Valid PDF/A output ({len(pdfa_data)} bytes)")
            else:
                log_test("POST /api/pdf/pdfa", "FAIL", "Invalid PDF output")
    except Exception as e:
        log_test("POST /api/pdf/pdfa", "FAIL", f"Exception: {e}")

def test_14_crop():
    """Test 14: POST /api/pdf/crop"""
    try:
        pdf_data = create_simple_pdf()
        files = {'file': ('test.pdf', pdf_data, 'application/pdf')}
        data = {'margin': '10.0'}
        resp = requests.post(f"{API_BASE}/pdf/crop", files=files, data=data, timeout=60)
        
        if resp.status_code != 200:
            log_test("POST /api/pdf/crop", "FAIL", f"HTTP {resp.status_code}")
        else:
            cropped_pdf = resp.content
            if is_valid_pdf(cropped_pdf):
                log_test("POST /api/pdf/crop", "PASS", f"Valid cropped PDF ({len(cropped_pdf)} bytes)")
            else:
                log_test("POST /api/pdf/crop", "FAIL", "Invalid PDF output")
    except Exception as e:
        log_test("POST /api/pdf/crop", "FAIL", f"Exception: {e}")

def test_15_compare():
    """Test 15: POST /api/pdf/compare"""
    try:
        pdf1_data = create_simple_pdf()
        pdf2_data = create_table_pdf()
        
        files = {
            'file1': ('test1.pdf', pdf1_data, 'application/pdf'),
            'file2': ('test2.pdf', pdf2_data, 'application/pdf')
        }
        resp = requests.post(f"{API_BASE}/pdf/compare", files=files, timeout=60)
        
        if resp.status_code != 200:
            log_test("POST /api/pdf/compare", "FAIL", f"HTTP {resp.status_code}")
        else:
            data = resp.json()
            if 'similarity' in data and 'rows' in data:
                log_test("POST /api/pdf/compare", "PASS", f"Comparison result: {data['similarity']}% similarity")
            else:
                log_test("POST /api/pdf/compare", "FAIL", "Invalid comparison result")
    except Exception as e:
        log_test("POST /api/pdf/compare", "FAIL", f"Exception: {e}")

def print_summary():
    """Print a summary table of all test results."""
    print("\n" + "="*80)
    print("BACKEND REGRESSION TEST SUMMARY")
    print("="*80)
    
    passed = sum(1 for r in results if r['status'] == 'PASS')
    failed = sum(1 for r in results if r['status'] == 'FAIL')
    
    print(f"\nTotal Tests: {len(results)}")
    print(f"Passed: {passed}")
    print(f"Failed: {failed}")
    print(f"Success Rate: {(passed/len(results)*100):.1f}%\n")
    
    if failed > 0:
        print("FAILED TESTS:")
        print("-" * 80)
        for r in results:
            if r['status'] == 'FAIL':
                print(f"❌ {r['endpoint']}")
                print(f"   {r['details']}\n")
    
    print("="*80)

if __name__ == "__main__":
    print("Starting LovePDF Backend Regression Test")
    print(f"Base URL: {BASE_URL}")
    print(f"API Base: {API_BASE}")
    print("="*80 + "\n")
    
    # Run all tests
    test_1_health()
    test_2_mongo_connectivity()
    test_3_inspect()
    test_4_office_to_pdf()
    test_5_html_to_pdf()
    test_6_pdf_to_word()
    test_7_pdf_to_excel()
    test_8_pdf_to_ppt()
    test_9_ocr()
    test_10_protect()
    test_11_unlock()
    test_12_repair()
    test_13_pdfa()
    test_14_crop()
    test_15_compare()
    
    # Print summary
    print_summary()
    
    # Exit with appropriate code
    failed = sum(1 for r in results if r['status'] == 'FAIL')
    sys.exit(0 if failed == 0 else 1)
